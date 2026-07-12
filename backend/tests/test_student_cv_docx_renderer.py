"""Tests for StudentCvDocxRenderer.

We construct StudentProfile / StudentProfileEntry ORM objects by hand
(no DB) and round-trip the resulting bytes through `docx.Document(...)`
to inspect the rendered output. Focus is on:

  * All five templates produce a valid file that python-docx can reopen.
  * Populated section headings appear exactly once.
  * Empty sections don't produce any heading.
  * The student's name and per-entry titles land in the doc.
  * Unknown template slug falls back to `classic` gracefully.
"""
from __future__ import annotations

from datetime import UTC, date, datetime
from decimal import Decimal
from io import BytesIO
from uuid import uuid4

from docx import Document

from app.application.services.student_cv_docx_renderer import (
    StudentCvDocxRenderer,
)
from app.application.services.student_cv_renderer import StudentCvRenderer
from app.infrastructure.db.models.student_profile import (
    StudentProfile,
    StudentProfileEntry,
)


def _profile() -> StudentProfile:
    p = StudentProfile()
    p.user_id = uuid4()
    p.full_name = "Amina El-Sayed"
    p.professional_email = "amina@example.com"
    p.phone = "+20 100 555 0100"
    p.location = "Cairo, Egypt"
    p.date_of_birth = date(2002, 5, 14)
    p.college = "Cairo University"
    p.department = "Faculty of Engineering"
    p.degree = "BSc"
    p.major = "Computer Science"
    p.graduation_year = 2026
    p.gpa = Decimal("3.75")
    p.headline = "Software engineering intern seeker"
    p.summary = (
        "Third-year CS student with a strong interest in backend "
        "systems and machine learning."
    )
    p.links = {
        "linkedin": "https://linkedin.com/in/amina",
        "github": "https://github.com/amina",
        "portfolio": "https://amina.dev",
    }
    p.interests = []
    p.photo_offset_x = 50
    p.photo_offset_y = 50
    p.photo_zoom = 100
    p.completed_steps = []
    p.current_step = "preview"
    p.cv_template_slug = "classic"
    return p


def _entry(
    *,
    kind: str,
    title: str,
    organization: str | None = None,
    description: str | None = None,
    details: dict | None = None,
    sort_order: int = 0,
    is_current: bool = False,
    start: date | None = None,
    end: date | None = None,
    url: str | None = None,
) -> StudentProfileEntry:
    e = StudentProfileEntry()
    e.id = uuid4()
    e.user_id = uuid4()
    e.kind = kind
    e.title = title
    e.organization = organization
    e.start_date = start
    e.end_date = end
    e.is_current = is_current
    e.description = description
    e.url = url
    e.details = details or {}
    e.sort_order = sort_order
    e.created_at = datetime.now(UTC)
    e.updated_at = datetime.now(UTC)
    return e


def _full_entries() -> list[StudentProfileEntry]:
    return [
        _entry(kind="skill", title="Python"),
        _entry(kind="skill", title="TypeScript"),
        _entry(kind="skill", title="PostgreSQL"),
        _entry(
            kind="language",
            title="Arabic",
            details={"proficiency": "Native"},
        ),
        _entry(
            kind="language",
            title="English",
            details={"proficiency": "Fluent"},
        ),
        _entry(
            kind="project",
            title="Careero — job matcher",
            description="Matches students to internships from a résumé.",
            details={
                "roles": ["backend", "database"],
                "tech_stack": ["Python", "FastAPI", "PostgreSQL"],
                "features": "search, apply flow, notifications",
                "hardest_part": "ranking under sparse feedback",
            },
            start=date(2025, 1, 1),
            is_current=True,
            url="https://github.com/amina/careero",
        ),
        _entry(
            kind="project",
            title="LinkAlyze",
            description="Static analyzer for TypeScript module boundaries.",
            details={
                "roles": ["frontend", "testing"],
                "tech_stack": ["TypeScript", "React"],
            },
        ),
        _entry(kind="course", title="Data Structures and Algorithms"),
        _entry(kind="course", title="Machine Learning"),
        _entry(
            kind="certificate",
            title="AWS Certified Cloud Practitioner",
            organization="Amazon Web Services",
            end=date(2025, 3, 1),
        ),
        _entry(
            kind="volunteer",
            title="Coding Club Mentor",
            organization="Cairo Uni Coding Society",
            description="Ran weekly workshops for freshmen on Python basics.",
            start=date(2024, 9, 1),
            is_current=True,
        ),
        # No 'award' or 'extracurricular' — those sections must NOT appear.
    ]


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _para_full_text(para) -> str:
    """`Paragraph.text` only walks direct `<w:r>` children — it skips
    text inside `<w:hyperlink>`. Walk every `<w:t>` in the paragraph
    so we pick up link labels ("LinkedIn", "GitHub", "Portfolio") too.
    """
    return "".join(t.text or "" for t in para._p.iter(f"{_W_NS}t"))


def _paragraph_texts(bytes_: bytes) -> list[str]:
    """All paragraph texts across the document, including any inside
    table cells and any `<w:hyperlink>` runs (link display labels)."""
    doc = Document(BytesIO(bytes_))
    texts = [_para_full_text(p) for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(_para_full_text(p))
    return texts


def _hyperlink_targets(bytes_: bytes) -> list[str]:
    """Every external hyperlink URL in the document (in the order the
    relationships were declared). One entry per unique URL."""
    doc = Document(BytesIO(bytes_))
    part = doc.part
    return [
        rel.target_ref
        for rel in part.rels.values()
        if rel.reltype.endswith("/hyperlink") and rel.is_external
    ]


def test_all_templates_render_valid_docx() -> None:
    p = _profile()
    entries = _full_entries()
    renderer = StudentCvDocxRenderer()
    for slug in ("classic", "modern", "minimal", "academic", "creative"):
        b = renderer.render_docx(
            profile=p, entries=entries, template_slug=slug
        )
        assert isinstance(b, bytes)
        assert len(b) > 2000  # a real .docx is at least a few KB
        # Round-trip parse succeeds and the student's name appears.
        texts = _paragraph_texts(b)
        assert any(p.full_name in t for t in texts), f"missing name in {slug}"


def test_populated_sections_appear_once() -> None:
    p = _profile()
    entries = _full_entries()
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=entries, template_slug="classic"
    )
    texts = _paragraph_texts(b)

    def _count(needle: str) -> int:
        return sum(1 for t in texts if t.strip() == needle)

    for heading in (
        "Professional Summary",
        "Education",
        "Skills",
        "Languages",
        "Projects",
        "Relevant Coursework",
        "Certificates",
        "Volunteer Experience",
    ):
        assert _count(heading) == 1, f"heading {heading!r} should appear once"

    # Sections with no entries — no heading rendered.
    for absent in ("Awards", "Extracurriculars"):
        assert _count(absent) == 0, f"heading {absent!r} should not appear"


def test_project_titles_and_narrative_render() -> None:
    p = _profile()
    entries = _full_entries()
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=entries, template_slug="modern"
    )
    texts = _paragraph_texts(b)
    joined = "\n".join(texts)

    for title in ("Careero — job matcher", "LinkAlyze"):
        assert title in joined
    # Project narrative composed from structured details.
    assert "FastAPI" in joined
    assert "backend" in joined or "backend development" in joined


def test_partial_profile_only_renders_populated_sections() -> None:
    """Student who has only filled Basics + Education gets a CV with
    just their name/contact/education — no empty section headings."""
    p = _profile()
    p.summary = None
    p.headline = None
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=[], template_slug="classic"
    )
    texts = _paragraph_texts(b)
    text = "\n".join(texts)
    assert p.full_name in text
    assert "Education" in text  # education block renders from profile fields
    for absent in (
        "Professional Summary",
        "Skills",
        "Projects",
        "Certificates",
        "Volunteer Experience",
        "Awards",
    ):
        assert absent not in text


def test_unknown_template_falls_back_to_classic() -> None:
    p = _profile()
    entries = _full_entries()
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=entries, template_slug="bogus-slug"
    )
    # Renders successfully — fallback path is the same as if we'd
    # explicitly asked for classic.
    texts = _paragraph_texts(b)
    assert any(p.full_name in t for t in texts)
    assert any(t.strip() == "Skills" for t in texts)


def test_missing_profile_still_renders() -> None:
    """Even with profile=None the renderer returns a valid file — that
    matches the PDF path (which shows the "Your name" placeholder)."""
    b = StudentCvDocxRenderer().render_docx(
        profile=None, entries=[], template_slug="classic"
    )
    assert isinstance(b, bytes) and len(b) > 2000
    texts = _paragraph_texts(b)
    assert any("Your name" in t for t in texts)


def test_modern_uses_two_column_layout() -> None:
    """Modern's dark sidebar carries the visual identity across from
    the PDF — it's built as a one-row, two-cell table. Skills and
    languages live in the sidebar; the rest lives in the main cell."""
    p = _profile()
    entries = _full_entries()
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=entries, template_slug="modern"
    )
    doc = Document(BytesIO(b))
    assert len(doc.tables) == 1, "modern should have exactly one layout table"
    table = doc.tables[0]
    assert len(table.rows) == 1 and len(table.columns) == 2
    side_texts = [par.text for par in table.rows[0].cells[0].paragraphs]
    main_texts = [par.text for par in table.rows[0].cells[1].paragraphs]
    joined_side = "\n".join(side_texts)
    joined_main = "\n".join(main_texts)
    # Skills / Languages in sidebar
    assert "Skills" in joined_side
    assert "Languages" in joined_side
    # Name + project headings in main
    assert p.full_name in joined_main
    assert "Projects" in joined_main
    # And Skills should NOT appear again in main
    assert "Skills" not in joined_main


def test_link_line_only_includes_populated_keys() -> None:
    p = _profile()
    p.links = {"linkedin": "https://linkedin.com/in/amina"}  # no github/portfolio
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=[], template_slug="classic"
    )
    text = "\n".join(_paragraph_texts(b))
    # Display label appears; the raw URL is only stored as a hyperlink
    # target, not printed in the doc.
    assert "LinkedIn" in text
    assert "GitHub" not in text
    assert "Portfolio" not in text
    # Hyperlink resolves to the correct URL.
    targets = _hyperlink_targets(b)
    assert "https://linkedin.com/in/amina" in targets
    assert not any("github" in t for t in targets)


def test_links_are_real_hyperlinks_with_short_labels() -> None:
    """Every populated link shows a short display label ('LinkedIn',
    'GitHub', 'Portfolio') and drops the URL into a Word hyperlink
    relationship — not as plain text."""
    p = _profile()
    entries = _full_entries()
    for slug in ("classic", "modern", "minimal", "academic", "creative"):
        b = StudentCvDocxRenderer().render_docx(
            profile=p, entries=entries, template_slug=slug
        )
        text = "\n".join(_paragraph_texts(b))
        for label in ("LinkedIn", "GitHub", "Portfolio"):
            assert label in text, f"{slug}: missing {label!r} display label"
        # Raw profile URLs must NOT appear as plain text on the page.
        assert "linkedin.com/in/amina" not in text, f"{slug}: URL leaked as text"
        assert "https://amina.dev" not in text, f"{slug}: URL leaked as text"
        # But they must exist as hyperlink relationship targets.
        targets = _hyperlink_targets(b)
        assert "https://linkedin.com/in/amina" in targets
        assert "https://github.com/amina" in targets
        assert "https://amina.dev" in targets


def test_html_preview_links_open_outside_sandboxed_iframe() -> None:
    html = StudentCvRenderer().render_html(
        profile=_profile(),
        entries=_full_entries(),
        template_slug="classic",
    )
    assert 'target="_blank" rel="noopener" href="https://github.com/amina"' in html
    assert 'target="_blank" rel="noopener" href="https://amina.dev"' in html


def test_unsafe_legacy_links_are_not_exported() -> None:
    profile = _profile()
    profile.links = {
        "portfolio": "file://attacker/share",
        "website": "javascript:alert(1)",
    }
    entry = _entry(
        kind="project",
        title="Unsafe project",
        url="smb://attacker/share",
    )

    context = StudentCvRenderer().build_context(profile=profile, entries=[entry])
    assert context["profile"]["links"] == {}
    assert context["sections"][0]["entries"][0]["url"] is None

    output = StudentCvDocxRenderer().render_docx(
        profile=profile,
        entries=[entry],
        template_slug="classic",
    )
    document = Document(BytesIO(output))
    targets = [rel.target_ref for rel in document.part.rels.values() if rel.is_external]
    assert not any(target.startswith(("file:", "javascript:", "smb:")) for target in targets)


def test_internship_ai_bullets_render_verbatim() -> None:
    """Internships store the LLM-polished summary+bullets on
    entry.details. The DOCX must render each bullet on its own line
    (not squash them into a paragraph) so ATS parsers pick them up."""
    p = _profile()
    entries = [
        _entry(
            kind="internship",
            title="Software Development Intern",
            organization="TechNova Solutions",
            details={
                "field": "software_engineering",
                "work_mode": "remote",
                "responsibilities": "helped test the website, fixed small bugs",
                "tools": ["GitHub", "VS Code"],
                "skills_gained": ["debugging", "collaboration"],
                "ai_summary": (
                    "Completed a software engineering internship "
                    "focused on website testing and bug reporting."
                ),
                "ai_bullets": [
                    "Assisted in testing website features and validating flows.",
                    "Reported bugs and supported the engineering team.",
                    "Contributed to minor frontend fixes under supervision.",
                ],
            },
            start=date(2024, 5, 1),
            end=date(2024, 8, 1),
        ),
    ]
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=entries, template_slug="classic"
    )
    texts = _paragraph_texts(b)
    joined = "\n".join(texts)
    assert "Internships" in joined
    assert "TechNova Solutions" in joined
    for bullet in (
        "Assisted in testing website features",
        "Reported bugs and supported",
        "Contributed to minor frontend fixes",
    ):
        assert any(bullet in t for t in texts), f"missing bullet: {bullet!r}"


def test_internship_deterministic_fallback_when_no_ai_bullets() -> None:
    """A student who saved an internship without asking the coach
    should still get usable bullets on their CV — composed from
    responsibilities/achievements/tools by the deterministic helper."""
    from app.application.services.student_cv_renderer import (
        _deterministic_internship_bullets,
    )
    from app.infrastructure.db.models.student_profile import (
        StudentProfileEntry,
    )
    e = StudentProfileEntry()
    e.kind = "internship"
    e.title = "Marketing Intern"
    e.organization = "BrightWave Media"
    e.description = None
    e.details = {
        "responsibilities": "I helped with social media campaigns\nI created content ideas",
        "tools": ["Canva", "Meta Ads Manager"],
        "skills_gained": ["copywriting"],
    }
    bullets = _deterministic_internship_bullets(e)
    assert 1 <= len(bullets) <= 4
    # Weak "I helped" → strong "Supported"
    assert any(b.startswith("Supported") for b in bullets)
    # Tools bullet composes when raw fields don't already cover them
    joined = " ".join(bullets)
    # At least one bullet references the tools or skills
    assert (
        "Canva" in joined
        or "Meta Ads Manager" in joined
        or "copywriting" in joined
    )


def test_photo_fills_circle_on_landscape_and_portrait() -> None:
    """Regression: at wizard-default zoom (100), landscape *and*
    portrait photos both need to fully cover the circle. Verified by
    passing synthetic non-square images and checking that the circular
    PNG we produce has no fully-transparent pixel inside its inscribed
    circle (i.e. no visible white gaps)."""
    from PIL import Image

    from app.application.services.student_cv_docx_renderer import (
        _make_circular_photo,
    )

    def _has_gaps(png_bytes: bytes) -> bool:
        im = Image.open(BytesIO(png_bytes)).convert("RGBA")
        w, h = im.size
        cx, cy, r = w / 2, h / 2, min(w, h) / 2 - 2
        px = im.load()
        # Sample points along a smaller inner ring — must all be
        # opaque with a non-white pixel behind them.
        import math
        for i in range(24):
            a = 2 * math.pi * i / 24
            x = int(cx + math.cos(a) * (r * 0.6))
            y = int(cy + math.sin(a) * (r * 0.6))
            _, _, _, alpha = px[x, y]
            if alpha < 250:
                return True
        return False

    for W, H in ((1000, 600), (600, 1000), (2000, 400), (400, 2000), (500, 500)):
        raw_img = Image.new("RGB", (W, H), color=(30, 90, 200))
        raw_buf = BytesIO()
        raw_img.save(raw_buf, format="PNG")
        result = _make_circular_photo(
            raw_buf.getvalue(),
            offset_x=50, offset_y=50, zoom=100,
        )
        assert result is not None
        assert not _has_gaps(result), (
            f"{W}x{H} at zoom=100 should fully cover the circle"
        )


def test_headline_renders_right_after_name() -> None:
    """Every template puts the "message" (headline) directly after the
    student's name."""
    p = _profile()
    b = StudentCvDocxRenderer().render_docx(
        profile=p, entries=[], template_slug="classic"
    )
    texts = _paragraph_texts(b)
    # Find the name paragraph and check the headline appears within
    # the next couple of paragraphs.
    name_idx = next(
        i for i, t in enumerate(texts) if p.full_name in t
    )
    assert any(
        p.headline in texts[i]
        for i in range(name_idx + 1, min(name_idx + 4, len(texts)))
    ), "headline should follow the name paragraph"
