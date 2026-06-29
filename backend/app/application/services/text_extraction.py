"""Text extraction from PDF / DOCX / plain-text uploads.

Pure functions over bytes. The CV / LinkedIn ingest services use these to
turn the raw upload into text the LLM can structure.
"""
from __future__ import annotations

import io
import logging

from pdfminer.high_level import extract_text as _pdfminer_extract
from docx import Document as _DocxDocument

logger = logging.getLogger(__name__)


class TextExtractionError(RuntimeError):
    """Raised when a file cannot be parsed into text (image-only PDF, etc.)."""


PDF_TYPES = {"application/pdf", "application/x-pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TEXT_TYPES = {"text/plain", "text/markdown"}


def extract_text(*, content: bytes, content_type: str) -> str:
    """Dispatch on content_type. Returns extracted plain text or raises.

    PDF extraction uses pdfminer.six's high-level interface; the result is
    typically usable but loses bullet structure (which is fine — the LLM
    structuring step doesn't care).

    DOCX extraction concatenates paragraphs with newlines; tables are
    flattened into ``cell\\tcell`` lines.
    """
    ct = (content_type or "").split(";", 1)[0].strip().lower()
    if ct in PDF_TYPES:
        return _extract_pdf(content)
    if ct in DOCX_TYPES:
        return _extract_docx(content)
    if ct in TEXT_TYPES or ct.startswith("text/"):
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")
    raise TextExtractionError(f"Unsupported content type: {content_type!r}")


def _extract_pdf(content: bytes) -> str:
    try:
        text = _pdfminer_extract(io.BytesIO(content)) or ""
    except Exception as exc:  # pdfminer raises a grab-bag of exception classes
        raise TextExtractionError(f"PDF parse failed: {exc}") from exc
    cleaned = "\n".join(line.rstrip() for line in text.splitlines() if line.strip())
    if not cleaned:
        # Likely an image-only PDF — pdfminer returns "" for those.
        raise TextExtractionError(
            "No text extracted — this looks like an image-only or scanned PDF. "
            "Try DOCX or paste the content instead."
        )
    return cleaned


def _extract_docx(content: bytes) -> str:
    try:
        doc = _DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise TextExtractionError(f"DOCX parse failed: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)
    cleaned = "\n".join(parts)
    if not cleaned:
        raise TextExtractionError("DOCX appears to be empty.")
    return cleaned
