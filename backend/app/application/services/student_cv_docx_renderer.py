"""StudentCvDocxRenderer — programmatic .docx CV builder.

Sibling of `StudentCvRenderer` (PDF/HTML). Reuses `build_context()` from
that class so the "what belongs where" logic is defined exactly once —
empty-section filtering, entry ordering, project narrative composition
all come from the same source of truth.

We deliberately generate DOCX **programmatically** with `python-docx`
rather than converting the PDF. Every template is Word-editable, uses
standard fonts (Calibri / Georgia), and avoids ATS-hostile constructs:

  * No text boxes, no floating shapes.
  * No image-based text — headings are real text runs.
  * No SmartArt or icons.

Link handling: URLs are rendered as clickable Word hyperlinks with a
short label ("LinkedIn", "GitHub", "Portfolio") rather than the raw
URL — nicer to read, still ATS-friendly (hyperlinks are plain text
runs plus an `r:id` reference).

Photo handling: cropped + circle-masked server-side via Pillow using
the wizard's `photo_offset_x/y` + `photo_zoom`. Uses the same CSS-like
"cover" positioning as the PDF (`background-position` + `background-size`)
so the DOCX photo shows exactly what the student cropped.

Layouts:
  * classic — mini header table (photo | name+headline+contact+links).
  * modern — full-width 2-column layout (dark navy sidebar | main).
  * creative — mini header table with purple shading (band + photo + text).
  * minimal / academic — single column.

Tables used for layout are single-row, borderless, plain — a well-
supported DOCX construct that Word, Google Docs, and modern ATS
parsers all handle. No text boxes, no floating shapes.
"""
from __future__ import annotations

import io
import logging
from collections.abc import Callable
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageOps

from app.application.services.student_cv_renderer import StudentCvRenderer
from app.application.url_policy import safe_external_http_url
from app.infrastructure.db.models.student_profile import (
    StudentProfile,
    StudentProfileEntry,
)

logger = logging.getLogger(__name__)


# --- Palette (matches the PDF templates) ---------------------------------

_INK = RGBColor(0x1F, 0x29, 0x37)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_CLASSIC_ACCENT = RGBColor(0x25, 0x63, 0xEB)
_CLASSIC_ACCENT_HEX = "2563EB"
_MODERN_DARK_HEX = "0F172A"
_MODERN_ACCENT = RGBColor(0x38, 0xBD, 0xF8)
_MODERN_ACCENT_HEX = "38BDF8"
_MODERN_SIDE_TEXT = RGBColor(0xE2, 0xE8, 0xF0)
_MINIMAL_INK = RGBColor(0x11, 0x18, 0x27)
_ACADEMIC_ACCENT = RGBColor(0x7C, 0x2D, 0x12)
_ACADEMIC_ACCENT_HEX = "7C2D12"
_ACADEMIC_META = RGBColor(0x57, 0x53, 0x4E)
_CREATIVE_PURPLE_HEX = "7C3AED"
_CREATIVE_ACCENT = RGBColor(0x7C, 0x3A, 0xED)


LINK_LABELS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("LinkedIn", ("linkedin",)),
    ("GitHub", ("github",)),
    ("Portfolio", ("portfolio", "website")),
)


def _link_pairs(links: dict[str, Any] | None) -> list[tuple[str, str]]:
    """Return ordered (label, url) pairs for the links that exist."""
    if not links:
        return []
    out: list[tuple[str, str]] = []
    for label, keys in LINK_LABELS:
        for k in keys:
            v = safe_external_http_url(links.get(k))
            if v:
                out.append((label, v))
                break
    return out


# --- Photo (crop + circular mask via Pillow) -----------------------------


def _make_circular_photo(
    raw: bytes,
    *,
    offset_x: int,
    offset_y: int,
    zoom: int,
    output_px: int = 320,
) -> bytes | None:
    """Reproduce the wizard's crop transform and mask the result to a
    circle. Returns PNG bytes with transparent corners.

    Positioning follows the CSS the PDF templates use
    (`background-position: X% Y%` + `background-size: <zoom>%`) with
    one important safeguard: we always scale the image up enough to
    fully cover the circle. The wizard's raw `background-size: 100%`
    can leave gaps on non-square photos (image width fills but height
    doesn't, or vice versa); the DOCX should never show those gaps.
    Effect: at defaults the photo is neatly centered inside the circle;
    the student's pan/zoom offsets still shift the visible slice.
    """
    try:
        img = Image.open(io.BytesIO(raw))
        # Phone cameras stamp orientation into EXIF instead of rotating
        # the pixels. Apply the rotation before we start cropping so a
        # portrait phone photo doesn't come out sideways.
        img = ImageOps.exif_transpose(img)
        img = img.convert("RGBA")
    except Exception:
        logger.warning("Could not decode student photo for DOCX; skipping.")
        return None

    W, H = img.size
    if W == 0 or H == 0:
        return None
    zoom_pct = max(zoom or 100, 100)
    # Wizard uses `background-size: <zoom>%`, i.e. width = container *
    # zoom/100. Reproduce that…
    wizard_scale = (output_px * zoom_pct / 100.0) / W
    # …but never below the scale that covers the circle on both axes.
    cover_scale = max(output_px / W, output_px / H)
    scale = max(wizard_scale, cover_scale)

    NW = max(output_px, round(W * scale))
    NH = max(output_px, round(H * scale))
    scaled = img.resize((NW, NH), Image.LANCZOS)

    # CSS-style position clamped so the image always covers the circle.
    pos_x = round((output_px - NW) * (offset_x or 50) / 100.0)
    pos_y = round((output_px - NH) * (offset_y or 50) / 100.0)
    pos_x = max(min(pos_x, 0), output_px - NW)
    pos_y = max(min(pos_y, 0), output_px - NH)

    canvas = Image.new("RGBA", (output_px, output_px), (255, 255, 255, 255))
    canvas.paste(scaled, (pos_x, pos_y))

    mask = Image.new("L", (output_px, output_px), 0)
    ImageDraw.Draw(mask).ellipse((0, 0, output_px - 1, output_px - 1), fill=255)
    out = Image.new("RGBA", (output_px, output_px), (255, 255, 255, 0))
    out.paste(canvas, (0, 0), mask)

    buf = io.BytesIO()
    out.save(buf, format="PNG")
    return buf.getvalue()


# --- Low-level docx helpers ---------------------------------------------


Container = DocxDocument | _Cell


def _add_picture(
    container: Container, bytes_: bytes, *, width_inches: float
) -> Paragraph | None:
    try:
        if isinstance(container, _Cell):
            para = container.add_paragraph()
            run = para.add_run()
            run.add_picture(io.BytesIO(bytes_), width=Inches(width_inches))
            return para
        else:
            container.add_picture(
                io.BytesIO(bytes_), width=Inches(width_inches)
            )
            return container.paragraphs[-1]
    except Exception:
        logger.warning("DOCX photo insertion failed; continuing without photo.")
        return None


def _set_document_defaults(
    doc: DocxDocument, *, font_name: str, font_size_pt: float
) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    normal.font.color.rgb = _INK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_page_margins(doc: DocxDocument, *, inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _pPr(para: Paragraph) -> OxmlElement:
    return para.paragraph_format.element.get_or_add_pPr()


def _add_border(
    para: Paragraph,
    *,
    edge: str,
    color_hex: str = "444444",
    size: int = 6,
    space: int = 1,
) -> None:
    pPr = _pPr(para)
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    tag = f"w:{edge}"
    existing = pBdr.find(qn(tag))
    if existing is not None:
        pBdr.remove(existing)
    b = OxmlElement(tag)
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color_hex)
    pBdr.append(b)


def _add_para_shading(para: Paragraph, *, fill_hex: str) -> None:
    pPr = _pPr(para)
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _shade_cell(cell: _Cell, *, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(
    cell: _Cell,
    *,
    top: int = 100,
    bottom: int = 100,
    left: int = 150,
    right: int = 150,
) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcPr.append(mar)


def _cell_vertical_align_center(cell: _Cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tcPr.append(v)


def _remove_table_borders(table) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


def _new_run(
    para: Paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size_pt: float | None = None,
    color: RGBColor | None = None,
    small_caps: bool = False,
    all_caps: bool = False,
    font_name: str | None = None,
) -> None:
    run = para.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if small_caps:
        run.font.small_caps = True
    if all_caps:
        run.font.all_caps = True
    if font_name is not None:
        run.font.name = font_name


def _add_hyperlink(
    para: Paragraph,
    url: str,
    text: str,
    *,
    color_hex: str = "0563C1",
    bold: bool = False,
    size_pt: float | None = None,
    underline: bool = True,
) -> None:
    """Real Word hyperlink: clickable display text linked to a URL.
    python-docx has no first-class API so we build the OOXML by hand —
    same pattern used across the ecosystem (docxtpl, docx-hyperlink).
    """
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    rPr.append(c)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
        rPr.append(sz)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def _spacing(
    para: Paragraph,
    *,
    before_pt: float | None = None,
    after_pt: float | None = None,
    line_spacing: float | None = None,
) -> None:
    pf = para.paragraph_format
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def _joined(items: list[str], sep: str = " · ") -> str:
    return sep.join(x for x in items if x)


def _add_link_line(
    para: Paragraph,
    links: dict[str, Any] | None,
    *,
    color_hex: str,
    size_pt: float,
    separator: str = "  ·  ",
) -> bool:
    """Append clickable link labels ("LinkedIn", "GitHub", "Portfolio")
    into `para`, separated by `separator`. Returns True if anything was
    written.
    """
    pairs = _link_pairs(links)
    if not pairs:
        return False
    for i, (label, url) in enumerate(pairs):
        if i:
            _new_run(para, separator, color=_MUTED, size_pt=size_pt)
        _add_hyperlink(
            para, url, label,
            color_hex=color_hex,
            size_pt=size_pt,
            underline=False,
        )
    return True


def _add_photo(
    container: Container,
    photo_bytes: bytes | None,
    profile: dict[str, Any],
    *,
    width_inches: float = 1.2,
    align: int = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    if not photo_bytes:
        return
    circular = _make_circular_photo(
        photo_bytes,
        offset_x=profile.get("photo_offset_x") or 50,
        offset_y=profile.get("photo_offset_y") or 50,
        zoom=profile.get("photo_zoom") or 100,
    )
    if not circular:
        return
    para = _add_picture(container, circular, width_inches=width_inches)
    if para is not None:
        para.alignment = align
        _spacing(para, before_pt=0, after_pt=6)


# --- Body section renderers (container-agnostic) ------------------------


def _render_summary(container: Container, profile: dict[str, Any], *, style: dict) -> None:
    """Long-form summary block (headline is rendered separately in the
    header, not here — this section is body-level 'Professional Summary'
    prose only)."""
    summary = (profile.get("summary") or "").strip()
    if not summary:
        return
    _h2(container, "Professional Summary", style)
    p = container.add_paragraph(summary)
    _spacing(p, before_pt=0, after_pt=6, line_spacing=1.15)


def _render_education(container: Container, profile: dict[str, Any], *, style: dict) -> None:
    college = (profile.get("college") or "").strip()
    degree = (profile.get("degree") or "").strip()
    major = (profile.get("major") or "").strip()
    department = (profile.get("department") or "").strip()
    grad_year = profile.get("graduation_year")
    gpa = (profile.get("gpa") or "").strip() if profile.get("gpa") else None
    if not any([college, degree, major, department, grad_year, gpa]):
        return
    _h2(container, "Education", style)
    p = container.add_paragraph()
    if college:
        _new_run(p, college, bold=True)
    if grad_year:
        _new_run(p, f"    Expected {grad_year}", italic=True, color=_MUTED)
    _spacing(p, before_pt=0, after_pt=1)
    detail = ", ".join(x for x in [degree, major, department] if x)
    if detail:
        p2 = container.add_paragraph(detail)
        _spacing(p2, before_pt=0, after_pt=1)
    if gpa:
        p3 = container.add_paragraph()
        _new_run(p3, f"GPA: {gpa}", italic=True, color=_MUTED)
        _spacing(p3, before_pt=0, after_pt=6)


def _render_section_skill(container: Container, entries: list[dict[str, Any]], *, style: dict) -> None:
    labels = [e["title"] for e in entries if e.get("title")]
    if not labels:
        return
    p = container.add_paragraph(", ".join(labels))
    _spacing(p, before_pt=0, after_pt=6, line_spacing=1.15)


def _render_section_language(container: Container, entries: list[dict[str, Any]], *, style: dict) -> None:
    parts: list[str] = []
    for e in entries:
        title = (e.get("title") or "").strip()
        details = e.get("details") if isinstance(e.get("details"), dict) else {}
        proficiency = (details or {}).get("proficiency")
        if title and proficiency:
            parts.append(f"{title} ({proficiency})")
        elif title:
            parts.append(title)
    if not parts:
        return
    p = container.add_paragraph(", ".join(parts))
    _spacing(p, before_pt=0, after_pt=6)


def _render_section_course(container: Container, entries: list[dict[str, Any]], *, style: dict) -> None:
    for e in entries:
        title = (e.get("title") or "").strip()
        if not title:
            continue
        p = container.add_paragraph(style="List Bullet")
        _new_run(p, title)
        _spacing(p, before_pt=0, after_pt=0)


def _render_section_certificate(container: Container, entries: list[dict[str, Any]], *, style: dict) -> None:
    for e in entries:
        title = (e.get("title") or "").strip()
        issuer = (e.get("organization") or "").strip()
        end = (e.get("end_date") or "").strip() if e.get("end_date") else None
        p = container.add_paragraph()
        _new_run(p, title, bold=True)
        tail = _joined([issuer, end], sep=" · ")
        if tail:
            _new_run(p, f"  — {tail}", color=_MUTED)
        _spacing(p, before_pt=0, after_pt=3)


def _render_section_entry_block(container: Container, entries: list[dict[str, Any]], *, style: dict) -> None:
    entry_border_hex = style.get("entry_left_border")
    for e in entries:
        title = (e.get("title") or "").strip()
        org = (e.get("organization") or "").strip()
        start = (e.get("start_date") or "").strip() if e.get("start_date") else None
        end = (
            "Present"
            if e.get("is_current")
            else ((e.get("end_date") or "").strip() if e.get("end_date") else None)
        )
        p_head = container.add_paragraph()
        if title:
            _new_run(p_head, title, bold=True)
        head_tail = _joined([org, _joined([start or "", end or ""], sep=" – ")], sep=" · ")
        if head_tail:
            _new_run(p_head, f"  — {head_tail}", italic=True, color=_MUTED)
        _spacing(p_head, before_pt=2, after_pt=1)
        if entry_border_hex:
            _add_border(p_head, edge="left", color_hex=entry_border_hex, size=8, space=6)
            p_head.paragraph_format.left_indent = Inches(0.12)
        body = (e.get("narrative") or e.get("description") or "").strip()
        if body:
            p_body = container.add_paragraph(body)
            _spacing(p_body, before_pt=0, after_pt=6, line_spacing=1.15)
            if entry_border_hex:
                _add_border(p_body, edge="left", color_hex=entry_border_hex, size=8, space=6)
                p_body.paragraph_format.left_indent = Inches(0.12)
        url = safe_external_http_url(e.get("url"))
        if url:
            p_url = container.add_paragraph()
            _add_hyperlink(p_url, url, url, color_hex=style.get("link_hex", "0563C1"), size_pt=10)
            _spacing(p_url, before_pt=0, after_pt=6)
            if entry_border_hex:
                _add_border(p_url, edge="left", color_hex=entry_border_hex, size=8, space=6)
                p_url.paragraph_format.left_indent = Inches(0.12)


def _render_section_internship(
    container: Container, entries: list[dict[str, Any]], *, style: dict
) -> None:
    """Internship blocks are structured differently from projects: each
    entry has a title/org/date header, an optional 1-line summary, and
    2–4 bullet points (either LLM-polished or deterministic).
    """
    for e in entries:
        title = (e.get("title") or "").strip()
        org = (e.get("organization") or "").strip()
        start = (e.get("start_date") or "").strip() if e.get("start_date") else None
        end = (
            "Present"
            if e.get("is_current")
            else ((e.get("end_date") or "").strip() if e.get("end_date") else None)
        )
        p_head = container.add_paragraph()
        if title:
            _new_run(p_head, title, bold=True)
        head_tail = _joined([org, _joined([start or "", end or ""], sep=" – ")], sep=" · ")
        if head_tail:
            _new_run(p_head, f"  — {head_tail}", italic=True, color=_MUTED)
        _spacing(p_head, before_pt=2, after_pt=1)

        summary = (e.get("summary") or "").strip()
        if summary:
            p_sum = container.add_paragraph()
            _new_run(p_sum, summary, italic=True)
            _spacing(p_sum, before_pt=0, after_pt=2, line_spacing=1.15)

        bullets = [
            str(b).strip()
            for b in (e.get("bullets") or [])
            if isinstance(b, str) and b.strip()
        ]
        for bullet in bullets:
            p = container.add_paragraph(style="List Bullet")
            _new_run(p, bullet)
            _spacing(p, before_pt=0, after_pt=1, line_spacing=1.15)

        url = (e.get("url") or "").strip()
        if url:
            p_url = container.add_paragraph()
            _new_run(p_url, url, italic=True, color=_MUTED)
            _spacing(p_url, before_pt=0, after_pt=4)


_SECTION_RENDERERS: dict[
    str, Callable[[Container, list[dict[str, Any]], Any], None]
] = {
    "skill": _render_section_skill,
    "language": _render_section_language,
    "project": _render_section_entry_block,
    "internship": _render_section_internship,
    "course": _render_section_course,
    "certificate": _render_section_certificate,
    "volunteer": _render_section_entry_block,
    "award": _render_section_entry_block,
    "extracurricular": _render_section_entry_block,
}


def _h2(container: Container, text: str, style: dict) -> None:
    p = container.add_paragraph()
    _new_run(
        p,
        text,
        bold=True,
        size_pt=style.get("h2_size", 12),
        color=style.get("h2_color", _INK),
        small_caps=style.get("h2_small_caps", False),
        all_caps=style.get("h2_all_caps", False),
    )
    _spacing(
        p,
        before_pt=style.get("h2_before", 10),
        after_pt=style.get("h2_after", 3),
    )
    if style.get("h2_border"):
        _add_border(p, edge="bottom", color_hex=style["h2_border"], size=6)


def _sidebar_h3(cell: _Cell, text: str) -> None:
    """Modern sidebar mini-heading — small letter-spaced label in cyan
    accent on the navy background."""
    p = cell.add_paragraph()
    _new_run(
        p,
        text,
        bold=True,
        size_pt=10,
        color=_MODERN_ACCENT,
        all_caps=True,
    )
    _spacing(p, before_pt=8, after_pt=2)


def _render_body_common(
    container: Container,
    ctx: dict,
    *,
    style: dict,
    include_sidebar_kinds: bool = True,
) -> None:
    profile = ctx["profile"]
    _render_summary(container, profile, style=style)
    _render_education(container, profile, style=style)
    for section in ctx["sections"]:
        if not include_sidebar_kinds and section["kind"] in {"skill", "language"}:
            continue
        _h2(container, section["label"], style)
        renderer = _SECTION_RENDERERS.get(section["kind"])
        if renderer:
            renderer(container, section["entries"], style=style)


def _identity_block(
    container: Container,
    profile: dict[str, Any],
    *,
    name_size: float,
    name_color: RGBColor,
    name_all_caps: bool,
    name_align: int,
    headline_color: RGBColor,
    contact_color: RGBColor,
    contact_size: float,
    link_color_hex: str,
    text_align: int,
) -> None:
    """Name → headline (message) → contact line → links line.
    Every template's identity column uses this — only the paint varies.
    """
    # Name
    p_name = container.add_paragraph()
    p_name.alignment = name_align
    _new_run(
        p_name,
        profile["full_name"],
        bold=True,
        size_pt=name_size,
        all_caps=name_all_caps,
        color=name_color,
    )
    _spacing(p_name, before_pt=0, after_pt=2)

    # Headline (the "message" right after the name)
    headline = (profile.get("headline") or "").strip()
    if headline:
        p_head = container.add_paragraph()
        p_head.alignment = text_align
        _new_run(p_head, headline, italic=True, color=headline_color, size_pt=12)
        _spacing(p_head, before_pt=0, after_pt=4)

    # Contact
    contact = _joined([
        profile.get("professional_email") or "",
        profile.get("phone") or "",
        profile.get("location") or "",
        profile.get("date_of_birth") or "",
    ])
    if contact:
        p_c = container.add_paragraph()
        p_c.alignment = text_align
        _new_run(p_c, contact, color=contact_color, size_pt=contact_size)
        _spacing(p_c, before_pt=0, after_pt=1)

    # Links (clickable labels)
    links = profile.get("links") or {}
    pairs = _link_pairs(links)
    if pairs:
        p_l = container.add_paragraph()
        p_l.alignment = text_align
        _add_link_line(
            p_l, links,
            color_hex=link_color_hex,
            size_pt=contact_size,
        )
        _spacing(p_l, before_pt=0, after_pt=4)


# --- Per-template renderers ---------------------------------------------


def _render_classic(
    ctx: dict, photo_bytes: bytes | None, photo_mime: str | None
) -> DocxDocument:
    """Header: mini 2-cell table (photo | identity). Body: single
    column with a blue accent."""
    doc = Document()
    _set_document_defaults(doc, font_name="Calibri", font_size_pt=11)
    _set_page_margins(doc, inches=0.8)
    profile = ctx["profile"]

    has_photo = bool(photo_bytes)
    if has_photo:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _remove_table_borders(table)
        photo_w = Inches(1.4)
        text_w = Inches(5.4)
        table.columns[0].width = photo_w
        table.columns[1].width = text_w
        row = table.rows[0]
        row.cells[0].width = photo_w
        row.cells[1].width = text_w
        _set_cell_margins(row.cells[0], top=0, bottom=0, left=0, right=200)
        _set_cell_margins(row.cells[1], top=40, bottom=40, left=0, right=0)
        _cell_vertical_align_center(row.cells[1])
        _add_photo(row.cells[0], photo_bytes, profile, width_inches=1.25, align=WD_ALIGN_PARAGRAPH.LEFT)
        _identity_block(
            row.cells[1],
            profile,
            name_size=22,
            name_color=_INK,
            name_all_caps=True,
            name_align=WD_ALIGN_PARAGRAPH.LEFT,
            headline_color=_MUTED,
            contact_color=_MUTED,
            contact_size=10,
            link_color_hex=_CLASSIC_ACCENT_HEX,
            text_align=WD_ALIGN_PARAGRAPH.LEFT,
        )
    else:
        _identity_block(
            doc,
            profile,
            name_size=24,
            name_color=_INK,
            name_all_caps=True,
            name_align=WD_ALIGN_PARAGRAPH.CENTER,
            headline_color=_MUTED,
            contact_color=_MUTED,
            contact_size=10,
            link_color_hex=_CLASSIC_ACCENT_HEX,
            text_align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Thin rule under the whole header block.
    rule = doc.add_paragraph()
    _spacing(rule, before_pt=2, after_pt=2)
    _add_border(rule, edge="bottom", color_hex="D1D5DB", size=6, space=6)

    _render_body_common(
        doc,
        ctx,
        style={
            "h2_size": 12,
            "h2_color": _CLASSIC_ACCENT,
            "h2_before": 10,
            "h2_after": 3,
            "h2_border": "E5E7EB",
            "link_hex": _CLASSIC_ACCENT_HEX,
        },
    )
    return doc


def _render_modern(
    ctx: dict, photo_bytes: bytes | None, photo_mime: str | None
) -> DocxDocument:
    """Two-column layout: dark navy sidebar (photo + contact + links +
    skills + languages) on the left; main column (name + headline +
    summary + education + everything else) on the right. Both cells
    are inside a single-row borderless table.
    """
    doc = Document()
    _set_document_defaults(doc, font_name="Calibri", font_size_pt=11)
    _set_page_margins(doc, inches=0.5)
    profile = ctx["profile"]

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _remove_table_borders(table)
    sidebar_w = Inches(2.5)
    main_w = Inches(4.9)
    table.columns[0].width = sidebar_w
    table.columns[1].width = main_w
    row = table.rows[0]
    row.cells[0].width = sidebar_w
    row.cells[1].width = main_w

    side = row.cells[0]
    main = row.cells[1]

    _shade_cell(side, fill_hex=_MODERN_DARK_HEX)
    _set_cell_margins(side, top=280, bottom=280, left=220, right=220)
    _set_cell_margins(main, top=140, bottom=140, left=260, right=200)

    # Sidebar
    _add_photo(side, photo_bytes, profile, width_inches=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    contact_bits = [
        profile.get("professional_email") or "",
        profile.get("phone") or "",
        profile.get("location") or "",
    ]
    if any(contact_bits):
        _sidebar_h3(side, "Contact")
        for line in contact_bits:
            if not line:
                continue
            p = side.add_paragraph()
            _new_run(p, line, color=_MODERN_SIDE_TEXT, size_pt=9.5)
            _spacing(p, before_pt=0, after_pt=1)

    pairs = _link_pairs(profile.get("links"))
    if pairs:
        _sidebar_h3(side, "Links")
        for label, url in pairs:
            p = side.add_paragraph()
            _add_hyperlink(
                p, url, label,
                color_hex=_MODERN_ACCENT_HEX,
                size_pt=9.5,
                bold=True,
                underline=False,
            )
            _spacing(p, before_pt=0, after_pt=1)

    for section in ctx["sections"]:
        if section["kind"] == "skill":
            _sidebar_h3(side, section["label"])
            titles = [e["title"] for e in section["entries"] if e.get("title")]
            if titles:
                p = side.add_paragraph()
                _new_run(p, ", ".join(titles), color=_MODERN_SIDE_TEXT, size_pt=9.5)
                _spacing(p, before_pt=0, after_pt=4, line_spacing=1.15)
        elif section["kind"] == "language":
            _sidebar_h3(side, section["label"])
            for e in section["entries"]:
                title = (e.get("title") or "").strip()
                details = e.get("details") if isinstance(e.get("details"), dict) else {}
                prof = (details or {}).get("proficiency")
                line = f"{title} — {prof}" if title and prof else title
                if not line:
                    continue
                p = side.add_paragraph()
                _new_run(p, line, color=_MODERN_SIDE_TEXT, size_pt=9.5)
                _spacing(p, before_pt=0, after_pt=1)

    # Main column: reuse the shared identity block for name + headline.
    # First paragraph in the cell already exists — the identity block
    # will append its own paragraphs after it, so drop the shell one.
    if main.paragraphs and not main.paragraphs[0].runs:
        p0 = main.paragraphs[0]
        p0._element.getparent().remove(p0._element)

    p_name = main.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _new_run(p_name, profile["full_name"], bold=True, size_pt=26, color=_INK)
    _spacing(p_name, before_pt=0, after_pt=2)
    headline = (profile.get("headline") or "").strip()
    if headline:
        p_h = main.add_paragraph()
        _new_run(p_h, headline, italic=True, color=_MUTED, size_pt=12)
        _spacing(p_h, before_pt=0, after_pt=8)

    _render_body_common(
        main,
        ctx,
        style={
            "h2_size": 11,
            "h2_color": _MODERN_ACCENT,
            "h2_small_caps": True,
            "h2_before": 10,
            "h2_after": 3,
            "h2_border": "1E293B",
            "link_hex": _MODERN_ACCENT_HEX,
        },
        include_sidebar_kinds=False,
    )
    return doc


def _render_minimal(
    ctx: dict, photo_bytes: bytes | None, photo_mime: str | None
) -> DocxDocument:
    doc = Document()
    _set_document_defaults(doc, font_name="Calibri", font_size_pt=10.5)
    _set_page_margins(doc, inches=0.6)
    _identity_block(
        doc,
        ctx["profile"],
        name_size=22,
        name_color=_MINIMAL_INK,
        name_all_caps=False,
        name_align=WD_ALIGN_PARAGRAPH.LEFT,
        headline_color=_MUTED,
        contact_color=_MUTED,
        contact_size=10,
        link_color_hex="111827",
        text_align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    rule = doc.add_paragraph()
    _spacing(rule, before_pt=0, after_pt=2)
    _add_border(rule, edge="bottom", color_hex="111827", size=6, space=6)

    _render_body_common(
        doc,
        ctx,
        style={
            "h2_size": 10,
            "h2_color": _MINIMAL_INK,
            "h2_small_caps": True,
            "h2_before": 8,
            "h2_after": 2,
            "link_hex": "111827",
        },
    )
    return doc


def _render_academic(
    ctx: dict, photo_bytes: bytes | None, photo_mime: str | None
) -> DocxDocument:
    doc = Document()
    _set_document_defaults(doc, font_name="Georgia", font_size_pt=11)
    _set_page_margins(doc, inches=0.8)
    _identity_block(
        doc,
        ctx["profile"],
        name_size=24,
        name_color=_INK,
        name_all_caps=False,
        name_align=WD_ALIGN_PARAGRAPH.CENTER,
        headline_color=_ACADEMIC_META,
        contact_color=_ACADEMIC_META,
        contact_size=10,
        link_color_hex=_ACADEMIC_ACCENT_HEX,
        text_align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    rule = doc.add_paragraph()
    _spacing(rule, before_pt=2, after_pt=2)
    _add_border(rule, edge="bottom", color_hex="A8A29E", size=6, space=6)

    _render_body_common(
        doc,
        ctx,
        style={
            "h2_size": 12,
            "h2_color": _ACADEMIC_ACCENT,
            "h2_before": 10,
            "h2_after": 3,
            "h2_border": "A8A29E",
            "link_hex": _ACADEMIC_ACCENT_HEX,
        },
    )
    return doc


def _render_creative(
    ctx: dict, photo_bytes: bytes | None, photo_mime: str | None
) -> DocxDocument:
    """Purple 'band' header: mini 2-cell table (photo | identity), both
    cells shaded purple so it reads as a single colored banner. Body
    stays clean below.
    """
    doc = Document()
    _set_document_defaults(doc, font_name="Calibri", font_size_pt=11)
    _set_page_margins(doc, inches=0.6)
    profile = ctx["profile"]

    has_photo = bool(photo_bytes)
    if has_photo:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _remove_table_borders(table)
        photo_w = Inches(1.6)
        text_w = Inches(5.6)
        table.columns[0].width = photo_w
        table.columns[1].width = text_w
        row = table.rows[0]
        row.cells[0].width = photo_w
        row.cells[1].width = text_w
        photo_cell = row.cells[0]
        text_cell = row.cells[1]
        _shade_cell(photo_cell, fill_hex=_CREATIVE_PURPLE_HEX)
        _shade_cell(text_cell, fill_hex=_CREATIVE_PURPLE_HEX)
        _set_cell_margins(photo_cell, top=180, bottom=180, left=200, right=100)
        _set_cell_margins(text_cell, top=200, bottom=200, left=100, right=200)
        _cell_vertical_align_center(text_cell)
        _add_photo(photo_cell, photo_bytes, profile, width_inches=1.4, align=WD_ALIGN_PARAGRAPH.CENTER)
        _identity_block(
            text_cell,
            profile,
            name_size=26,
            name_color=_WHITE,
            name_all_caps=False,
            name_align=WD_ALIGN_PARAGRAPH.LEFT,
            headline_color=_WHITE,
            contact_color=_WHITE,
            contact_size=10.5,
            link_color_hex="FFFFFF",
            text_align=WD_ALIGN_PARAGRAPH.LEFT,
        )
    else:
        # No-photo variant: everything in a single shaded paragraph
        # block. Name / headline / contact / links stay on the purple
        # band.
        p_name = doc.add_paragraph()
        _new_run(p_name, profile["full_name"], bold=True, size_pt=28, color=_WHITE)
        _add_para_shading(p_name, fill_hex=_CREATIVE_PURPLE_HEX)
        p_name.paragraph_format.left_indent = Inches(0.2)
        p_name.paragraph_format.right_indent = Inches(0.2)
        _spacing(p_name, before_pt=6, after_pt=2)
        headline = (profile.get("headline") or "").strip()
        if headline:
            p = doc.add_paragraph()
            _new_run(p, headline, italic=True, color=_WHITE, size_pt=12)
            _add_para_shading(p, fill_hex=_CREATIVE_PURPLE_HEX)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            _spacing(p, before_pt=0, after_pt=4)
        contact = _joined([
            profile.get("professional_email") or "",
            profile.get("phone") or "",
            profile.get("location") or "",
        ])
        if contact:
            p = doc.add_paragraph()
            _new_run(p, contact, color=_WHITE, size_pt=10.5)
            _add_para_shading(p, fill_hex=_CREATIVE_PURPLE_HEX)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            _spacing(p, before_pt=0, after_pt=2)
        pairs = _link_pairs(profile.get("links"))
        if pairs:
            p = doc.add_paragraph()
            _add_link_line(
                p, profile.get("links"),
                color_hex="FFFFFF",
                size_pt=10.5,
            )
            _add_para_shading(p, fill_hex=_CREATIVE_PURPLE_HEX)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            _spacing(p, before_pt=0, after_pt=6)

    _render_body_common(
        doc,
        ctx,
        style={
            "h2_size": 12,
            "h2_color": _CREATIVE_ACCENT,
            "h2_before": 12,
            "h2_after": 3,
            "entry_left_border": "EDE9FE",
            "link_hex": _CREATIVE_PURPLE_HEX,
        },
    )
    return doc


_DOCX_TEMPLATE_REGISTRY: dict[
    str, Callable[[dict, bytes | None, str | None], DocxDocument]
] = {
    "classic": _render_classic,
    "modern": _render_modern,
    "minimal": _render_minimal,
    "academic": _render_academic,
    "creative": _render_creative,
}
_DEFAULT_DOCX_SLUG = "classic"


class StudentCvDocxRenderer:
    """Programmatic .docx CV builder. Wraps
    `StudentCvRenderer.build_context` so section grouping /
    empty-section skipping / project narrative composition come from
    exactly one place across both PDF and DOCX outputs.
    """

    def __init__(self) -> None:
        self._context_builder = StudentCvRenderer()

    @staticmethod
    def list_template_slugs() -> set[str]:
        return set(_DOCX_TEMPLATE_REGISTRY)

    def _resolve_builder(
        self, slug: str | None
    ) -> Callable[[dict, bytes | None, str | None], DocxDocument]:
        if slug and slug in _DOCX_TEMPLATE_REGISTRY:
            return _DOCX_TEMPLATE_REGISTRY[slug]
        return _DOCX_TEMPLATE_REGISTRY[_DEFAULT_DOCX_SLUG]

    def render_docx(
        self,
        *,
        profile: StudentProfile | None,
        entries: list[StudentProfileEntry],
        photo_bytes: bytes | None = None,
        photo_mime: str | None = None,
        template_slug: str | None = None,
    ) -> bytes:
        ctx = self._context_builder.build_context(
            profile=profile,
            entries=entries,
            photo_bytes=photo_bytes,
            photo_mime=photo_mime,
        )
        builder = self._resolve_builder(template_slug)
        doc = builder(ctx, photo_bytes, photo_mime)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
